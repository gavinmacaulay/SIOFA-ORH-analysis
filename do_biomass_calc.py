# -*- coding: utf-8 -*-
"""
Created on Sat Oct 30 12:29:30 2021

@author: gavin
"""
#%%
from pathlib import Path
import matplotlib.pyplot as plt
from matplotlib.offsetbox import AnchoredText
import pandas as pd
from docx import Document
from haversine import haversine, Unit
import math
import numpy as np

import sys
sys.path.append(r'C:\Users\gavin\Data - not synced\Code\python\lsss')
from read_lsss_db_export import read_lsss_db_export

projectDir = Path(r'C:\Users\gavin\OneDrive - Aqualyd Limited\Documents\Aqualyd\Projects\2021-05 SIO ORH survey analysis')
resultsDir = projectDir.joinpath('Results')

lsssReportDir = projectDir.joinpath(r'Data\LSSS_DATA\S10000_PWill Watch[1]\Reports')

exportFile = lsssReportDir.joinpath('ListUserFile20__L0.0-3397.1.xml')

metaDataFile = projectDir.joinpath(r'Data\Metadata.xlsx')

# Read in the echo-integrated data
acoCat = 1 # 'other'
acoCat = 5842 # orange roughy

integrals = read_lsss_db_export(20, exportFile, acoCat)

# Assign a feature and transect id to all sA values (then we can select particular features, etc).

# Get the transect/survey metadata
md = pd.read_excel(metaDataFile, sheet_name='Transects', parse_dates=['start_time', 'end_time']).fillna('')

# for each entry in integrals, find the transect it belongs to using the transect start and end times
features = []
areas = []
snapshots = []
transects = []
survey_ids = []
survey_types = []
for i, row in integrals.iterrows():
    match = md[(row.time >= md.start_time) & (row.time <= md.end_time)]
    if len(match) > 0:
        features.append(match.feature.array[0])
        areas.append(match.area.array[0])
        snapshots.append(str(match.snapshot.array[0]))
        transects.append(str(match.transect.array[0]))
        survey_ids.append(str(match.survey_id.array[0]))
        survey_types.append(match.survey_type.array[0])
    else:
        features.append('')
        areas.append('')
        snapshots.append('')
        transects.append('')
        survey_ids.append('')
        survey_types.append('')
        
integrals = integrals.assign(feature=features, area=areas, snapshot=snapshots, 
                             transect=transects, survey_id=survey_ids, survey_type=survey_types)

# drop all rows that don't have a feature, area, snapshot, transect, or survey_id as these
# are outside the survey. Since the code above ensures that this is a all or nothing
# we only need to select on one of these columns.
integrals = integrals.loc[integrals['feature'] != '']

# save integrals in csv to be loaded into a gis from which the survey boundaries
# can be drawn and the area calculated.
integrals.to_csv(projectDir.joinpath(r'Data/Metadata_for_survey_boundaries_generated.csv'), index=False)

# and on the assumption that the work in qgis has been done, load the areas in
survey_areas = pd.read_csv(projectDir.joinpath(r'Data/Survey_areas_exported_from_qgis.csv'), dtype={'survey_id': str})

# and load in the centre points of the surveys that are stars
star_centres = pd.read_csv(projectDir.joinpath(r'Data/feature star centres.csv'), dtype={'survey_id': str})
# and rename the position columns
star_centres = star_centres.rename(columns={"X": "lon", "Y": "lat"})

#%%
##########################################################
# do the biomass estimation

# storage of per-transect survey results
r_per_transect = []

# and storage of per-survey results
r_per_survey = []

# Mean fish length and weight to use, 
mean_orh_length = 44.6 # [cm]
w_m = 0.3348 * np.power(mean_orh_length, 2.3636) # [g]
w_f = 0.2267 * np.power(mean_orh_length, 2.4856) # [g]
mean_orh_weight = (w_m + w_f)  / 2.0 / 1000.0 # [kg] assuming a 1:1 male/female ratio

mean_orh_TS = 16.37 * np.log10(mean_orh_length) - 77.17
mean_orh_sigma_bs = np.power(10.0, mean_orh_TS/10.0)

print(f'Using a mean length of {mean_orh_length:.1f} cm')
print(f'Using a mean weight of {mean_orh_weight:.2f} kg')
print(f'Using a TS of {mean_orh_TS:.1f} dB re 1m^2')

# Pull together the information needed to estimate biomass per survey
# since survey_id is the thing that is unique to each survey, work off that
surveys = integrals.groupby('survey_id')
for survey_id, survey in surveys:

    # we have the data from a single survey. Calculate the average sA per
    # transect and store
    t = survey.groupby('transect')
    transects_in_survey = []
    survey_type = survey.survey_type.values[0]

    print(f'Processing survey {survey.feature.values[0]} on {survey.time.min().date()} of type {survey_type}')

    # find the area for this survey
    survey_area = survey_areas[survey_areas.survey_id == survey_id].area.values[0] # [m^2]

    if survey_type == 'parallel':
        for j, cell in t: # over each transect in the survey
            transect_length = cell.distance.sum() # [nmi]
            transects_in_survey.append({'area': cell.area.values[0], 
                                        'feature': cell.feature.values[0],
                                        'snapshot': cell.snapshot.values[0],
                                        'transect': cell.transect.values[0],
                                        'transect_length': transect_length,
                                        'transect_start_time': cell.time.values[0],
                                        'mean_sA': cell.sA_sum.mean(), # [m^/nmi^2]
                                        'transect_rho': cell.sA_sum.mean() / (4.0*np.pi * mean_orh_sigma_bs * 1852*1852), # [fish/m^2]
                                        'survey_id': survey_id})
        transects_in_survey = pd.DataFrame(transects_in_survey)

        # mean sAs in transect j of survey i
        rho_ij = transects_in_survey.transect_rho.to_numpy() # [fish/m^2]
        n_i = transects_in_survey.transect.shape[0] # number of transects in survey i
        L_ij = transects_in_survey.transect_length # lengths of transects j of survey i
    
        L_i = 1.0/n_i * L_ij.sum() # [nmi]
        w_ij = L_ij / L_i # per transect weighting, based on transect lengths
    
        rho_i = 1.0/n_i * np.sum(w_ij * rho_ij) # mean density per survey i [fish/m^2]
        biomass = rho_i * survey_area * mean_orh_weight / 1000.0 # [tonnes]

        # variance of mean density in survey i        
        var_rho_i = np.sum(np.square(w_ij) * np.square(rho_ij - rho_i)) / (n_i*(n_i-1.0))
        if rho_i > 0:
            cv = 100.0 * math.sqrt(var_rho_i) / rho_i
        else:
            cv = 999.0
                
    else: # star
        # Using the notation of Doonan et al, 2003.
        #
        # transects are index j, samples within transect are index i
        lat = star_centres[star_centres.survey_id == survey_id].lat.values[0]
        lon = star_centres[star_centres.survey_id == survey_id].lon.values[0]

        dist = lambda sA_point: haversine(sA_point, (lat, lon), unit=Unit.KILOMETERS)
        
        for i, cell in t: # over each transect in the survey
            transect_length = cell.distance.sum() # [nmi]
            # for the current transect, work out the distance to the centre point
            
            sA_pos = np.array(list(zip(cell.lat.to_numpy(), cell.lon.to_numpy())))

            r_ij = np.array([dist(i) for i in sA_pos]) # [km]

            # fish density from each sA value
            y_ij = cell.sA_sum.to_numpy() / (4.0*np.pi * mean_orh_sigma_bs * 1852*1852) # [fish/m^2]

            # calc weighted mean fish density for each transect
            w_j = np.sum(y_ij * (r_ij / r_ij.sum())) # [fish/m^2]

            transects_in_survey.append({'area': cell.area.values[0], 
                                        'feature': cell.feature.values[0],
                                        'snapshot': cell.snapshot.values[0],
                                        'transect': cell.transect.values[0],
                                        'transect_length': transect_length,
                                        'transect_start_time': cell.time.values[0],
                                        'mean_sA': cell.sA_sum.mean(), # not relevant for the star method [m^/nmi^2]
                                        'transect_rho': w_j, # [fish/m^2]
                                        'survey_id': survey_id})
        
        # Now calculate the survey biomass
        transects_in_survey = pd.DataFrame(transects_in_survey)
        w = transects_in_survey.transect_rho.mean() # [fish/m^2]
        rho_i = w # for code outside this loop, to be consistent with the parallel code
        biomass = w * survey_area * mean_orh_weight / 1000.0 # [tonnes]
        
        n = transects_in_survey.transect.shape[0] # number of transects in this survey
        var_rho_i = np.sum(np.square(transects_in_survey.transect_rho - w.mean())) / (n_i*(n_i-1.0))
        if w.mean() > 0:
            cv = 100.0 * math.sqrt(var_rho_i) / w.mean()
        else:
            cv = 999.0
        
    r_per_transect.append(transects_in_survey)
    r_per_survey.append({'survey_id': survey_id,
                         'survey_start_timestamp': survey.time.min(),
                         'survey_start_date': survey.time.min().date(),
                         'area': survey.area.values[0],
                         'feature': survey.feature.values[0],
                         'survey_type': survey_type,
                         'num_transects': pd.unique(survey.transect).shape[0],
                         'mean_rho': rho_i, # [fish/m^2]
                         'survey_area': survey_area, # [m^2]
                         'biomass': biomass, # [t]
                         'cv': cv}) # [%]

r_per_transect = pd.concat(r_per_transect)
r_per_survey = pd.DataFrame(r_per_survey)

#%%
#################################################################
# write the biomass results to a Word table and an Excel file (this was an 
# extra request from SIOFA after submission of the draft ToR 3 report)
tableFile = projectDir.joinpath('docs').joinpath('ToR 3 report').joinpath('ToR 3 biomass table.docx')
excelFile = tableFile.parent.joinpath('biomass per feature.xlsx')

# adjust r_per_survey to suit the Word table
t1_word = r_per_survey[['area', 'feature', 'survey_start_date', 'survey_type', 'mean_rho', 'biomass', 'cv', 'num_transects', 'survey_area']]
col_format =           ['v',    'v',       'v',                 'v',           'v:.2f',    'v:.0f',   'v:.0f','v',           'v/1e6:.1f']
units =                ['',     '',        '',                  '',            '[fish/m^2]','[t]',    '[%]', '',             '[km^2]']

# Discard surveys with biomass of 0.0
t1_word = t1_word[t1_word.biomass > 0.0]

# Sort for table
t1_word = t1_word.sort_values(by=['area', 'feature', 'survey_start_date'])

# save to Excel
t1_word.to_excel(excelFile, sheet_name='Biomass results', index=False)

def fstr(template):
    return eval(f"f'{template}'")

document = Document()
table = document.add_table(rows=t1_word.shape[0]+1, cols=t1_word.shape[1])
for i, column in enumerate(t1_word):
    prev_cell_str = ''
    table.cell(0, i).text = column.replace('_', ' ').capitalize() + ' ' + units[i]
    for row in range(t1_word.shape[0]):
        v = (t1_word[column].iloc[row])
        if column == 'biomass':
            v = np.round(v, decimals=-2) # round to nearest 100 kg
        cell_str = fstr('{' + col_format[i] + '}')
        if (column == 'area'): # drop area name when it is in multiple consecutive rows
            if (cell_str == prev_cell_str):
                cell_str = ''
            else:
                prev_cell_str = cell_str

        if (column == 'feature'): # drop feature name when it is in multiple consecutive rows
            if (cell_str == prev_cell_str):
                cell_str = ''
            else:
                prev_cell_str = cell_str
            
        table.cell(row+1, i).text = cell_str
document.save(tableFile)

#%%
# while we're here, create a table of the environment data for the surveys that we have
# a biomass for
enviro = pd.read_csv(metaDataFile.parent.joinpath('Metadata_enviroment_generated.csv'))

enviro = enviro.assign(area = r_per_survey.area, feature = r_per_survey.feature, survey_start_date = r_per_survey.survey_start_date)

# keep just the surveys with biomass
e = enviro[enviro.survey_id.isin(r_per_survey[r_per_survey.biomass > 0.0].survey_id.astype('int64').to_list())]

tableFile = projectDir.joinpath('docs').joinpath('ToR 3 report').joinpath('environment table.docx')

# adjust e to suit the Word table
t1_word    = e[['area', 'feature', 'survey_start_date', 'mean_t', 'mean_s', 'mean_c', 'mean_abs', 'distance', 'time']]
col_format =   ['v',    'v',       'v',                 'v:.1f',  'v:.1f',  'v:.1f',  'v:.1f',    'v:.0f',     'v:.0f']
units =        ['',     '',        '',                  '[degC]', '[PSU]',  '[m/s]',  '[dB/km]',  '[km]',     '[h]']

# Sort for table
t1_word = t1_word.sort_values(by=['area', 'feature', 'survey_start_date'])

def fstr(template):
    return eval(f"f'{template}'")

document = Document()
table = document.add_table(rows=t1_word.shape[0]+1, cols=t1_word.shape[1])
for i, column in enumerate(t1_word):
    prev_cell_str = ''
    table.cell(0, i).text = column.replace('_', ' ').capitalize() + ' ' + units[i]
    for row in range(t1_word.shape[0]):
        v = (t1_word[column].iloc[row])
        cell_str = fstr('{' + col_format[i] + '}')
        if (column == 'area'): # drop area name when it is in multiple consecutive rows
            if (cell_str == prev_cell_str):
                cell_str = ''
            else:
                prev_cell_str = cell_str

        if (column == 'feature'): # drop feature name when it is in multiple consecutive rows
            if (cell_str == prev_cell_str):
                cell_str = ''
            else:
                prev_cell_str = cell_str
            
        table.cell(row+1, i).text = cell_str
document.save(tableFile)

#%%
# some plots of the results

fig, ((ax0, ax1)) = plt.subplots(nrows=1, ncols=2)

ax0.plot(r.time, r.sA_sum, '.')
ax0.set_ylabel('sA summed [XX]')
ax0.set_xlabel('Time [UTC]')

ax1.plot(r.lon, r.lat, '.')
ax1.set_ylabel('Latitude')
ax1.set_xlabel('Longitude')

#ax0.add_artist(AnchoredText(f'Cast {platform_num}\nCycle {ctd_to_use.cycle_number}\n{dists[i_min]:.0f} km\n{time_diffs[i_min]:.0f} hours', 
#                            loc='upper left', frameon=False))
 
fig.savefig(resultsDir.joinpath('sA_plots').joinpath('test.png'), bbox_inches='tight', pad_inches=0.05)
plt.close() 
    
